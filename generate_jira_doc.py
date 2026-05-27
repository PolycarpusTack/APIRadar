"""Generate DEVELOPMENT_PLAN.docx for JIRA Initiative population."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_colored_heading(doc, text, level, rgb):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*rgb)
    return h


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


EPIC_COLOR = (31, 73, 125)    # dark blue
STORY_COLOR = (68, 114, 196)  # medium blue
TASK_COLOR = (0, 0, 0)


def build():
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("API Contract Radar Monitor — Development Plan", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "This document is structured for JIRA import. Each EPIC maps to a JIRA Epic, "
        "each Story maps to a JIRA Story, and each Task maps to a JIRA Sub-task."
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC A
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC A — Tracer Bullet: OpenAPI Diff CLI", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Thin end-to-end slice — `drift check` parses two OpenAPI YAML files and "
        "posts a PR comment listing Breaking Changes. No Blast Radius yet.\n"
        "Business value: Gives a platform team proof that the tool catches breaking changes "
        "before merge. Unblocks consumer work.\n"
        "Exit criteria:\n"
        "  • drift check runs in CI on one real producer repo\n"
        "  • PR comment lists at least one Breaking Change with field path, kind, and severity\n"
        "  • radar-desktop launches; radar-api sidecar starts with SQLite; radar-ui loads\n"
        "  • Exit code 0/1/2 semantics documented"
    )

    add_colored_heading(doc, "Story A-1 · Project Skeleton & CI", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: S\nSo that all subsequent tasks start from a clean, runnable baseline with CI already green.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-1-T1", "PREPARATORY", "Init Cargo workspace with radar-cli, radar-api, radar-scanner crates; drift-dashboard pnpm workspace"],
            ["A-1-T2", "PREPARATORY", "GitHub Actions CI: cargo test + clippy + pnpm lint + pnpm build"],
            ["A-1-T3", "PREPARATORY", "Docker Compose: postgres:16, radar-api, drift-dashboard for local dev"],
            ["A-1-T4", "PREPARATORY", "sqlx migrate setup: drift-db crate, initial empty migration, run-on-startup flag"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story A-2 · Spike — OpenAPI Parser Selection", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: S · Type: SPIKE (throwaway)\nSo that Story A-3 starts with a confirmed library choice, not a guess.")
    add_table(doc, ["Spike Question"],
        [
            ["Does `oas3` (Rust) handle $ref resolution and vendor extensions without panicking on real-world specs?"],
            ["Does `openapiv3` crate support OpenAPI 3.1 discriminators?"],
            ["What is parse time on a 500 KB spec with 300 paths?"],
        ], [7.0])

    add_colored_heading(doc, "Story A-3 · Parse OpenAPI YAML → Typed Diff", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: M\nSo that the tool produces a structured list of Changes I can inspect before any CI wiring.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-3-T1", "FEATURE", "SpecVersion struct + parse_openapi(path) with $ref resolution"],
            ["A-3-T2", "FEATURE", "diff(base, head) → Vec<Change> — field removal, addition, type change, required→optional"],
            ["A-3-T3", "FEATURE", "classify_severity(change) → Severity — breaking vs non-breaking rules per OpenAPI semantics"],
            ["A-3-T4", "FEATURE", "JSON and human-readable table output renderers"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story A-4 · `drift check` CLI Command", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: S\nSo that I can run a single command against two specs and get a coloured table on stdout.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-4-T1", "FEATURE", "drift check subcommand with --base, --head, --spec, --format, --json, --no-color flags"],
            ["A-4-T2", "FEATURE", "Terminal colour rendering: breaking=red, non-breaking-risky=amber, safe=teal; respects NO_COLOR"],
            ["A-4-T3", "FEATURE", "Exit code semantics: 0=clean, 1=breaking changes found, 2=parse/config error"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story A-5 · Policy File (.radar.yml)", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: S\nSo that teams can choose warn-only vs block without changing the CLI invocation.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-5-T1", "FEATURE", "Parse .radar.yml config; default values when file absent"],
            ["A-5-T2", "FEATURE", "Policy evaluation: block_on: never | any_break | active_consumers; exit code decision"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story A-6 · GitHub PR Comment", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: M\nSo that the reviewer sees the breaking changes inline without leaving GitHub.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-6-T1", "FEATURE", "Detect current PR number from GITHUB_SHA / GITHUB_REF env vars"],
            ["A-6-T2", "FEATURE", "POST/PATCH GitHub comment via REST API; idempotent (find-then-update)"],
            ["A-6-T3", "FEATURE", "Markdown comment template: summary header, changes table, blast-radius placeholder, policy verdict"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story A-7 · radar-api Stub — Diff Submission", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: S\nSo that diff results are persisted for future dashboard and blast-radius use.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-7-T1", "FEATURE", "axum route POST /v1/services/:id/diffs; service, diff, change tables via sqlx"],
            ["A-7-T2", "FEATURE", "Service token auth middleware (bearer token, env-seeded for P0)"],
            ["A-7-T3", "FEATURE", "GET /v1/services/:id/diffs list endpoint"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story A-8 · radar-ui + radar-desktop Electron Shell (SQLite mode)", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: M\nSo that I can open a desktop app, point it at a spec, and see diffs without configuring any infrastructure.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["A-8-T1", "PREPARATORY", "radar-ui pnpm workspace with Vite 6 + React 19 + Tailwind + shadcn/ui; radar-desktop with electron-vite"],
            ["A-8-T2", "FEATURE", "radar-api SQLite mode: --db sqlite:PATH flag; sqlx AnyDatabase feature; same migrations on SQLite"],
            ["A-8-T3", "FEATURE", "Electron main process: spawn radar-api child; wait for health-check; terminate on quit"],
            ["A-8-T4", "FEATURE", "Minimal radar-ui home screen: service list, Run Check button calling radar-api via fetch"],
            ["A-8-T5", "FEATURE", "electron-builder config: Windows NSIS installer, macOS DMG, Linux AppImage; GitHub release job"],
        ], [1.0, 1.2, 4.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC B
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC B — Consumer Blast Radius & Release Notes", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Consumer Registry + Usage Telemetry Ingest + Blast Radius on OpenAPI + Release Notes CLI\n"
        "Business value: Closes the loop — PR comment now names which services break and how recently they called the affected field.\n"
        "Exit criteria:\n"
        "  • Three consumers registered; PR comment names them with last-seen timestamps\n"
        "  • drift explain --release-notes outputs valid Markdown from a real Diff\n"
        "  • Policy block_on: active_consumers blocks CI correctly"
    )

    add_colored_heading(doc, "Story B-1 · Consumer Registry API", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: M\nSo that I declare my dependency on a producer once and get notified of future breaks automatically.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["B-1-T1", "FEATURE", "consumer + subscription tables, migrations"],
            ["B-1-T2", "FEATURE", "POST /v1/consumers, POST /v1/services/:id/subscriptions, GET /v1/services/:id/consumers routes"],
            ["B-1-T3", "FEATURE", "drift register CLI subcommand: reads .radar.yml, calls POST /v1/consumers + subscription"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story B-2 · Usage Event Ingest", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: M\nSo that the system knows which fields each consumer actually calls at runtime.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["B-2-T1", "FEATURE", "usage_event table (hypertable-ready schema); batch insert via COPY for throughput"],
            ["B-2-T2", "FEATURE", "POST /v1/usage/events — accepts array; rate-limit per service token"],
            ["B-2-T3", "FEATURE", "Retention cron job: delete events older than lookback_days × 3"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story B-3 · Blast Radius Computation", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: L\nSo that I see exactly which consumers called the field I'm removing, and when.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["B-3-T1", "FEATURE", "blast_radius(diff_id) → Vec<(Consumer, Confidence, LastSeen)> query"],
            ["B-3-T2", "FEATURE", "GET /v1/diffs/:id/blast-radius endpoint"],
            ["B-3-T3", "FEATURE", "Update drift check PR comment to include blast radius table"],
            ["B-3-T4", "FEATURE", "Policy block_on: active_consumers — evaluate blast radius and set exit code"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story B-4 · Release Notes Generator", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P1 · Size: M\nSo that I can paste one command's output into GitHub Releases and consumers know exactly what changed.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["B-4-T1", "FEATURE", "Fetch Diff + Blast Radius from radar-api; populate template structured sections deterministically"],
            ["B-4-T2", "FEATURE", "Claude API call for narrative sections (breaking changes plain-language + per-consumer one-liner); prompt-cached"],
            ["B-4-T3", "FEATURE", "--out FILE and --post-github-release output modes"],
        ], [1.0, 1.2, 4.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC C
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC C — Multi-format + Dashboard + Playground", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: GraphQL + protobuf parsers; tree-sitter static call-site scanner; dashboard v1; "
        "Interactive Playground (Scalar); pre-sales sandbox environment\n"
        "Exit criteria:\n"
        "  • drift check works on GraphQL SDL and protobuf inputs\n"
        "  • Dashboard shows cross-service trend view in both browser and Electron\n"
        "  • Playground tab shows Try It for any registered producer's spec\n"
        "  • PostgreSQL mode verified; web self-host confirmed"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Dependencies"],
        [
            ["C-1", "Spike — GraphQL schema diff library", "S", "—"],
            ["C-2", "Spike — protobuf / buf diff approach", "S", "—"],
            ["C-3", "GraphQL SDL parser + Diff", "M", "C-1"],
            ["C-4", "Protobuf proto3 parser + Diff", "M", "C-2"],
            ["C-5", "tree-sitter consumer repo scanner (TS/Python/Go)", "L", "B-1"],
            ["C-6", "tree-sitter Rust + Java grammars", "M", "C-5"],
            ["C-7", "call_site table + scanner job (cron)", "M", "C-5"],
            ["C-8", "Blast radius: union usage events + call sites", "S", "B-3, C-7"],
            ["C-9", "radar-ui full shell (nav, dark theme, React Router)", "M", "A-8"],
            ["C-10", "radar-ui: Diffs list + Diff detail with blast radius table", "M", "C-9"],
            ["C-11", "radar-ui: KPI cards (breaking-changes-30d, consumers-at-risk)", "S", "C-10"],
            ["C-12", "Scalar Playground integration — browser and Electron", "M", "C-9"],
            ["C-13", "Sandbox environment config (pre-sales base URL + auth injection)", "S", "C-12"],
            ["C-14", "PostgreSQL mode; Docker Compose for web self-host; migration parity test", "M", "A-8-T2"],
            ["C-15", "Web deployment: radar-api serves Vite bundle from /app; nginx config", "S", "C-14"],
            ["C-16", "Design system token audit across radar-ui + Electron window chrome", "S", "C-9–C-15"],
        ], [0.8, 3.2, 0.6, 2.4])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC D
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC D — Hardening, Policy Engine, SaaS-viable Deploy", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Migration guide generator; full policy engine; multi-org scale; GitHub Release automation; "
        "performance tests; security review; runbook\n"
        "Mode: HARDENING — no new features, only completion, verification, and documentation"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Dependencies"],
        [
            ["D-1", "Migration guide generator (per-consumer, Claude prose)", "M", "B-4"],
            ["D-2", "Full policy engine (allow_override_with: label:drift-ack)", "M", "A-5"],
            ["D-3", "--post-github-release automation for release notes", "S", "B-4"],
            ["D-4", "Multi-org: OIDC dashboard auth; org-scoped service tokens", "M", "A-7"],
            ["D-5", "TimescaleDB hypertable migration for usage_event (opt-in)", "S", "B-2"],
            ["D-6", "Performance test suite: check p95 < 5 s, blast-radius p95 < 300 ms", "M", "All"],
            ["D-7", "Security review: threat model verification, rate limits, token audit", "M", "All"],
            ["D-8", "End-to-end smoke test automation (Playwright on dashboard)", "M", "C-9+"],
            ["D-9", "Runbook: deploy, rollback, on-call procedures", "S", "All"],
            ["D-10", "Public OpenAPI spec at docs/openapi.yaml", "S", "A-7, B-1–B-3"],
            ["D-11", "SBOM (syft), cosign-signed release binaries, cargo audit", "S", "All"],
        ], [0.8, 3.6, 0.6, 2.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC E
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC E — Durable Evidence & Differentiator Hardening", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Normalize blast-radius evidence into append-only impact_evidence records; harden CLI fail-mode "
        "semantics; advance scanner to S2; prove the full differentiator flow with fixture-driven demo scenario\n"
        "Business value: Closes the evidence gap — blast radius is now backed by durable, explainable, "
        "expiry-aware Evidence records. PR comment shows exactly what is known and why CI is blocking.\n"
        "Exit criteria:\n"
        "  • E2E demo scenario test passes: field removed → billing-svc (high, runtime) + mobile-gateway (low, static) → block\n"
        "  • No blast-radius entry returned without at least one impact_evidence record\n"
        "  • CLI fail-open / fail-closed / warn behavior explicit, tested, written to Policy Decision\n"
        "  • PR comment evidence table renders correctly\n"
        "Status: ALL STORIES DONE ✓"
    )

    add_colored_heading(doc, "Story E-1 · impact_evidence Table and Blast-Radius Normalization", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: L · Status: DONE\nSo that every consumer listed in blast radius has at least one traceable, timestamped Evidence record.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-1-T1", "PREPARATORY", "Migration 007_impact_evidence.sql — impact_evidence table with all fields"],
            ["E-1-T2", "FEATURE", "blast_radius() writer — produces impact_evidence rows; source_type based on evidence source"],
            ["E-1-T3", "FEATURE", "GET /v1/diffs/:id/blast-radius reader — reads from impact_evidence; supports ?max_age_days="],
            ["E-1-T4", "FEATURE", "Evidence expiry job — scheduled task deletes rows past expires_at"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story E-2 · Org-Scoped Authorization Audit", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: M · Status: DONE\nSo that no organization can read, enumerate, or modify another organization's data through any API endpoint.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-2-T1", "FEATURE", "Integration test suite: for each resource type, assert org A token returns 403 on org B resource IDs"],
            ["E-2-T2", "FEATURE", "Fix any org_id enforcement gaps discovered by E-2-T1; update middleware or query filters"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story E-3 · CLI Fail-Mode Hardening", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P0 · Size: S · Status: DONE\nSo that I can explicitly choose how the CI gate behaves when Radar API is unreachable, recorded in a Policy Decision.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-3-T1", "FEATURE", "fail_mode field in .radar.yml; parsing, validation, default=closed; all three mode behaviors"],
            ["E-3-T2", "FEATURE", "Policy Decision persistence — POST /v1/policy-decisions; drift check writes a record after every run"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story E-4 · PR Comment Evidence Rendering", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P1 · Size: M · Status: DONE\nSo that I can see exactly which consumers are at risk, why Radar believes that, and the policy verdict — all without leaving GitHub.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-4-T1", "FEATURE", "Evidence table renderer — Markdown table; sorts by confidence descending; truncates at 10 rows"],
            ["E-4-T2", "FEATURE", "Policy verdict section renderer — verdict badge, fail_mode, required action, override instruction"],
            ["E-4-T3", "FEATURE", "Update drift check PR comment to include evidence table and policy verdict sections"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story E-5 · Operation-Aware TypeScript Scanner (S2)", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P1 · Size: L · Status: DONE\nSo that blast-radius confidence reflects whether Radar knows which API operation a call site is targeting.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-5-T1", "FEATURE", "TypeScript generated-client detection — tree-sitter identifies API client method calls; emits (method_name → operation) mapping"],
            ["E-5-T2", "FEATURE", "Operation correlation logic — resolves detected method names against service spec operations"],
            ["E-5-T3", "FEATURE", "Confidence propagation — medium if operation populated, low if operation is NULL"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story E-6 · Demo Scenario Fixtures", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P1 · Size: M · Status: DONE\nSo that we can prove the field removed → evidence → block flow in a single repeatable test.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-6-T1", "PREPARATORY", "Create fixture directories: demo-payments-api (v1/v2 specs), demo-billing-svc (usage events), demo-mobile-gateway (TS client)"],
            ["E-6-T2", "FEATURE", "Expected PR comment fixture file; deterministic enough for byte-level assertion on structured sections"],
            ["E-6-T3", "FEATURE", "Integration test tests/demo_scenario.rs: loads fixtures, seeds DB, runs diff + blast radius + PR comment render"],
        ], [1.0, 1.2, 4.8])

    add_colored_heading(doc, "Story E-7 · Collection File Scanner — Postman Collection v2.1", 2, STORY_COLOR)
    doc.add_paragraph("Priority: P1 · Size: M · Status: DONE\nSo that Radar derives Consumer evidence from committed collection files without runtime telemetry.")
    add_table(doc, ["Task ID", "Hat", "Goal"],
        [
            ["E-7-T1", "PREPARATORY", "CollectionFile struct + parse_collection(path) — deserialise Postman v2.1 JSON; extract name, requests, URL, method, test scripts"],
            ["E-7-T2", "FEATURE", "Field-path extraction from test scripts — scan exec lines for .json().<field> patterns"],
            ["E-7-T3", "FEATURE", "Consumer auto-registration — upsert Consumer row using info.name; idempotent on (org_id, name)"],
            ["E-7-T4", "FEATURE", "Evidence writer — source_type=collection_file, confidence=medium; dedup on INSERT OR IGNORE"],
            ["E-7-T5", "FEATURE", "Scanner configuration — collection_paths glob list in scanner config TOML"],
        ], [1.0, 1.2, 4.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC F
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC F — Enterprise Workflow Packaging", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: GitHub Action; policy decisions table; acknowledgement workflow; "
        "Backstage and CODEOWNERS catalog importers; dashboard enterprise pages\n"
        "Exit criteria:\n"
        "  • New repo can install radar-action from docs in under 15 minutes\n"
        "  • PR comment clearly explains pass / warn / block with evidence\n"
        "  • Overrides recorded in acknowledgement table and visible in audit trail\n"
        "Status: ALL STORIES DONE ✓"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Status"],
        [
            ["F-1", "radar-action — GitHub Action composite action (TypeScript)", "L", "DONE"],
            ["F-2", "policy_decision table + persistence in radar-api", "M", "DONE (pre-delivered by E-3-T2)"],
            ["F-3", "Acknowledgement workflow — table + API endpoints + UI", "L", "DONE"],
            ["F-4", "Backstage catalog-info.yaml importer (polling job)", "M", "DONE"],
            ["F-5", "CODEOWNERS fallback importer", "S", "DONE"],
            ["F-6", "Catalog sync status in dashboard UI", "S", "DONE"],
            ["F-7", "Acknowledgement workflow in dashboard UI (diff detail page, ack button)", "M", "DONE"],
            ["F-8", "Audit trail page in dashboard UI (paginated, org-scoped)", "S", "DONE"],
            ["F-9", "Documentation — getting-started, backstage-integration, policy-reference, oidc-setup", "M", "DONE"],
        ], [0.8, 3.6, 0.6, 2.0])

    add_colored_heading(doc, "EPIC F+ — Evolution Rules", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Operator-defined severity overrides per change kind; glob path matching; "
        "server-side evaluation in diff response; CLI management; dashboard UI\n"
        "Status: ALL STORIES DONE ✓"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Status"],
        [
            ["F+-1", "Migration 016 + evolution_rule CRUD API (POST/GET/DELETE/PATCH)", "S", "DONE"],
            ["F+-2", "Server-side rule evaluator in get_diff — severity override + applied_rule field", "M", "DONE"],
            ["F+-3", "CLI radar rule subcommands (add, list, delete, toggle, test)", "S", "DONE"],
            ["F+-4", "Dashboard UI — EvolutionRulesPage + Governance nav section", "S", "DONE"],
        ], [0.8, 3.6, 0.6, 2.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC G
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC G — Runtime Evidence Collection", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: OTel collector processor; API gateway adapters; Node/Express and FastAPI middleware SDKs; "
        "evidence freshness dashboard; sampling controls; privacy documentation\n"
        "Exit criteria:\n"
        "  • At least one real service produces usage Evidence via OTel processor or gateway adapter without custom application code\n"
        "  • Dashboard shows evidence coverage by service and Consumer\n"
        "  • Stale evidence is visible and expires predictably\n"
        "Status: ALL STORIES DONE ✓"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Status"],
        [
            ["G-1", "Spike — OTel collector processor architecture", "S", "DONE (OTLP-over-HTTP in radar-api; no separate Go process)"],
            ["G-2", "OTel collector processor — OTLP JSON trace receiver in radar-api", "L", "DONE"],
            ["G-3", "API gateway adapter — Kong / NGINX log ingestion", "M", "DONE"],
            ["G-4", "Node/Express middleware SDK (@radar-monitor/sdk)", "M", "DONE"],
            ["G-5", "FastAPI middleware SDK (Python / ASGI)", "M", "DONE"],
            ["G-6", "Ingestion sampling controls (per-service sample rate, field-path allow/block list)", "S", "DONE"],
            ["G-7", "Evidence freshness dashboard page (coverage by service and Consumer)", "M", "DONE"],
            ["G-8", "Privacy/redaction documentation (runtime-usage-ingestion.md, security-and-privacy.md)", "S", "DONE"],
        ], [0.8, 3.6, 0.6, 2.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC H
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC H — Impact-Targeted Artifacts", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Diff+evidence-scoped test generation; deterministic templates per change kind; "
        "per-Consumer migration guides; release-note state workflow; generated artifacts in PR comment and dashboard\n"
        "Exit criteria:\n"
        "  • For each Breaking Change kind in the 5 templates, Radar generates at least one relevant test Artifact\n"
        "  • Release Notes include affected Consumers and Evidence\n"
        "  • Migration Guide is scoped to Consumer usage (call sites + runtime Evidence)\n"
        "Status: ALL STORIES DONE ✓"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Status"],
        [
            ["H-1", "Test generation from diff + evidence context (diff_id-only path, no Jira)", "L", "DONE"],
            ["H-2", "Deterministic test templates per change kind (5 templates: field_removed, required_changed, enum_value_removed, operation_removed, type_changed)", "M", "DONE"],
            ["H-3", "Per-Consumer migration guides scoped to Consumer usage and call sites", "M", "DONE"],
            ["H-4", "Release-note state workflow (draft → reviewed → published → superseded)", "M", "DONE"],
            ["H-5", "Generated test artifacts linked in PR comment", "S", "DONE"],
            ["H-6", "Artifact review/publish controls in dashboard UI (ReleaseNotesPage status transitions)", "M", "DONE"],
        ], [0.8, 3.6, 0.6, 2.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # EPIC I
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "EPIC I — Public Readiness", 1, EPIC_COLOR)
    doc.add_paragraph(
        "Theme: Polished demo repo; public documentation; self-host install guide; "
        "benchmark suite; SBOM and signed binaries; demo video script\n"
        "Mode: HARDENING — no new features\n"
        "Exit criteria:\n"
        "  • Public docs state the impact-aware contract drift product promise without caveats\n"
        "  • Demo works from clean clone with docker compose up\n"
        "  • CI is green; Enterprise pilot checklist complete\n"
        "Status: ALL STORIES DONE ✓"
    )
    add_table(doc, ["Story ID", "Title", "Size", "Status"],
        [
            ["I-1", "Demo repository set: fixtures/demo-payments-api/, demo-billing-svc/, demo-mobile-gateway/ with seeded runtime usage and GitHub workflow", "M", "DONE"],
            ["I-2", "Polished README with installation, demo scenario, and architecture diagram", "M", "DONE"],
            ["I-3", "docs/evidence-confidence.md, docs/security-and-privacy.md, docs/demo-scenario.md, docs/enterprise-deployment.md", "M", "DONE"],
            ["I-4", "Self-host install guide (docs/enterprise-deployment.md) — Docker Compose + PostgreSQL + OIDC", "S", "DONE (merged into I-3)"],
            ["I-5", "Benchmark suite: drift check p95 < 10 s, blast-radius p95 < 2 s, usage ingest p95 < 500 ms", "M", "DONE"],
            ["I-6", "SBOM (syft), cosign-signed release binaries, cargo audit, licensing review", "S", "DONE"],
            ["I-7", "docs/generated-artifacts.md and demo video script", "S", "DONE"],
        ], [0.8, 4.0, 0.6, 1.6])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # Global DoD
    # ══════════════════════════════════════════════════════════════════════════
    add_colored_heading(doc, "Global Definition of Done", 1, (80, 80, 80))
    doc.add_paragraph(
        "Every task must pass before the story is considered complete:\n"
        "  • Tests written first (TDD: failing test → implementation → green → refactor)\n"
        "  • No secrets in code or logs\n"
        "  • cargo clippy -- -D warnings passes / pnpm lint passes\n"
        "  • ≥ 80% line coverage on new code\n"
        "  • Contract tests written for any new public interface\n"
        "  • Feature flag present if task introduces a Claude API call or destructive migration\n"
        "  • Hand-off artifact written (updated Architecture Memory or Contract Snapshot)\n"
        "  • Domain Glossary consistent\n"
        "  • No duplicated logic — abstraction check passed"
    )

    out = r"C:\Projects\Ideas\21-API-Contract-Drift-Monitor\API-Contract-Radar-DEVELOPMENT-PLAN.docx"
    doc.save(out)
    print(f"Saved: {out}")


build()
